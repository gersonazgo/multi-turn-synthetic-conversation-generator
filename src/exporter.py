from __future__ import annotations

import json
from datetime import datetime, timezone, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .models import Conversation


STOP_REASON_LABELS = {
    "turns_ended": "Turn limit reached without success",
    "assistant_succeeded": "Assistant succeeded",
    "llm_transient_error": "ERROR — Persistent transient error (rate limit / timeout / server)",
    "llm_content_error": "ERROR — Empty or refused content from LLM",
    "llm_non_transient_error": "ERROR — Non-transient LLM failure",
}


def _format_timestamp(ts: str | None) -> str:
    if not ts:
        return "—"
    dt = datetime.fromisoformat(ts).astimezone(timezone.utc)
    return dt.strftime("%Y-%m-%d %H:%M:%S")


def save_batch(conversations: list[Conversation], output_dir: str) -> Path:
    path = Path(output_dir)
    path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    docx_path = path / f"batch_{timestamp}.docx"

    doc = Document()

    counts: dict[str, int] = {}
    for i, conversation in enumerate(conversations):
        if i > 0:
            doc.add_page_break()
        counts[conversation.test_case] = counts.get(conversation.test_case, 0) + 1
        seq = counts[conversation.test_case]
        label = f"Test Case: {conversation.test_case} #{seq:02d}"
        _append_conversation_to_doc(doc, conversation, label=label)

    doc.save(str(docx_path))
    return docx_path


def save_conversation(conversation: Conversation, output_dir: str) -> tuple[Path, Path]:
    path = Path(output_dir)
    path.mkdir(parents=True, exist_ok=True)

    base = conversation.id

    # JSON
    json_path = path / f"{base}.json"
    json_path.write_text(
        json.dumps(conversation.model_dump(), indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    # DOCX
    docx_path = path / f"{base}.docx"
    _save_docx(conversation, docx_path)

    return json_path, docx_path


def _add_cell_field(cell, label: str, value: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)


def _remove_cell_default_paragraph(cell) -> None:
    """Remove the default empty paragraph that python-docx creates in each cell."""
    first_p = cell.paragraphs[0]
    if not first_p.text:
        p_element = first_p._element
        p_element.getparent().remove(p_element)


def _append_conversation_to_doc(doc: Document, conversation: Conversation, label: str | None = None) -> None:
    heading = label or f"Test Case: {conversation.test_case}"
    title = doc.add_heading(heading, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = conversation.metadata
    stop_label = STOP_REASON_LABELS.get(
        conversation.conversation_stop_reason, conversation.conversation_stop_reason
    )

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    for cell in table.rows[0].cells:
        cell._element.get_or_add_tcPr().append(
            parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                      '<w:top w:val="none"/><w:left w:val="none"/>'
                      '<w:bottom w:val="none"/><w:right w:val="none"/>'
                      '</w:tcBorders>')
        )

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    _remove_cell_default_paragraph(left_cell)
    _remove_cell_default_paragraph(right_cell)

    # Left column — scenario data
    _add_cell_field(left_cell, "Capability", conversation.capability)
    _add_cell_field(left_cell, "Test Case", conversation.test_case)
    _add_cell_field(left_cell, "Scenario", conversation.scenario)
    _add_cell_field(left_cell, "Persona", conversation.persona)
    _add_cell_field(left_cell, "Collaboration", conversation.collaboration)

    # Right column — technical data
    _add_cell_field(right_cell, "Assistant Model", meta.assistant_model)
    _add_cell_field(right_cell, "Assistant Temperature", str(meta.assistant_temperature))
    _add_cell_field(right_cell, "User Model", meta.user_model)
    _add_cell_field(right_cell, "User Temperature", str(meta.user_temperature))
    _add_cell_field(right_cell, "Total Turns", str(meta.total_turns))
    _add_cell_field(right_cell, "Result", stop_label)
    _add_cell_field(right_cell, "Started", _format_timestamp(meta.started_at))
    _add_cell_field(right_cell, "Finished", _format_timestamp(meta.finished_at))

    # Separator
    doc.add_paragraph("─" * 50)

    # Messages
    for msg in conversation.messages:
        speaker = conversation.persona if msg.role == "user" else "Assistant"
        color = RGBColor(0x1A, 0x73, 0xE8) if msg.role == "user" else RGBColor(0x1E, 0x88, 0x50)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        run_name = p.add_run(f"{speaker}: ")
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_name.font.color.rgb = color

        run_content = p.add_run(msg.content)
        run_content.font.size = Pt(11)


def _save_docx(conversation: Conversation, filepath: Path) -> None:
    doc = Document()
    _append_conversation_to_doc(doc, conversation)
    doc.save(str(filepath))
